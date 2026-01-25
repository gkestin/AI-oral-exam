"""
RAG (Retrieval-Augmented Generation) Handler
=============================================
Handles document storage, retrieval, and context injection for personalized exams.
"""

import asyncio
import logging
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False
    np = None
from pydantic import BaseModel
import google.generativeai as genai
try:
    from google.cloud import storage
    STORAGE_AVAILABLE = True
except ImportError:
    STORAGE_AVAILABLE = False
    storage = None
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None
import tiktoken

from ..config import get_settings
from ..models import Assignment, User
from .firebase import get_firestore_service

logger = logging.getLogger(__name__)


class Document(BaseModel):
    """Represents a document for RAG."""
    id: str
    course_id: str
    assignment_id: Optional[str] = None
    student_id: Optional[str] = None
    type: str  # 'syllabus', 'project_report', 'slides', 'submission'
    title: str
    content: str
    metadata: Dict[str, Any] = {}
    embedding: Optional[List[float]] = None
    chunk_index: int = 0
    total_chunks: int = 1
    created_at: datetime = None


class RAGConfig(BaseModel):
    """Configuration for RAG operations."""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    max_context_length: int = 4000
    similarity_threshold: float = 0.7
    top_k: int = 5
    embedding_model: str = "models/text-embedding-004"


class RAGHandler:
    """Handles document processing and retrieval for contextualized questioning."""

    def __init__(self):
        settings = get_settings()
        genai.configure(api_key=settings.google_api_key)

        # Initialize embedding model
        self.embed_model = genai.GenerativeModel(model_name="models/text-embedding-004")

        # Token counter for text splitting
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

        # Storage client for document uploads
        self.storage_client = storage.Client() if (STORAGE_AVAILABLE and settings.gcp_project) else None
        self.bucket_name = "ai-oral-exam-documents"

        # In-memory vector store (in production, use Pinecone/Weaviate/Chroma)
        self.vector_store: Dict[str, List[Document]] = {}

        self.config = RAGConfig()

    def _generate_doc_id(self, content: str, metadata: Dict) -> str:
        """Generate unique document ID."""
        hash_input = f"{content[:100]}{json.dumps(metadata, sort_keys=True)}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.tokenizer.encode(text))

    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap."""
        sentences = text.replace('\n', '. ').split('. ')
        chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            sentence_tokens = self._count_tokens(sentence)

            if current_tokens + sentence_tokens > self.config.chunk_size:
                if current_chunk:
                    chunks.append('. '.join(current_chunk) + '.')
                    # Keep last few sentences for overlap
                    overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                    current_chunk = overlap_sentences + [sentence]
                    current_tokens = sum(self._count_tokens(s) for s in current_chunk)
                else:
                    # Single sentence exceeds chunk size
                    chunks.append(sentence + '.')
                    current_chunk = []
                    current_tokens = 0
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens

        if current_chunk:
            chunks.append('. '.join(current_chunk) + '.')

        return chunks

    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using Gemini."""
        try:
            result = genai.embed_content(
                model="models/text-embedding-004",
                content=text,
                task_type="retrieval_document"
            )
            return result['embedding']
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return [0.0] * 768  # Return zero vector on error

    def _calculate_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors."""
        if not NUMPY_AVAILABLE:
            # Fallback to pure Python implementation
            import math
            dot_product = sum(a * b for a, b in zip(vec1, vec2))
            norm1 = math.sqrt(sum(a * a for a in vec1))
            norm2 = math.sqrt(sum(b * b for b in vec2))

            if norm1 == 0 or norm2 == 0:
                return 0.0
            return dot_product / (norm1 * norm2)

        vec1 = np.array(vec1)
        vec2 = np.array(vec2)

        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)

        if norm1 == 0 or norm2 == 0:
            return 0.0

        return dot_product / (norm1 * norm2)

    async def process_document(
        self,
        file_path: str,
        course_id: str,
        doc_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Process a document and store its chunks with embeddings."""

        # Extract text based on file type
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self._extract_docx_text(file_path)
        elif file_ext == '.md':
            text = self._extract_markdown_text(file_path)
        elif file_ext in ['.txt', '.py', '.js', '.html']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Split into chunks
        chunks = self._split_text(text)

        # Create document objects with embeddings
        documents = []
        for i, chunk in enumerate(chunks):
            embedding = await self._generate_embedding(chunk)

            doc = Document(
                id=self._generate_doc_id(chunk, metadata or {}),
                course_id=course_id,
                assignment_id=metadata.get('assignment_id') if metadata else None,
                student_id=metadata.get('student_id') if metadata else None,
                type=doc_type,
                title=Path(file_path).stem,
                content=chunk,
                metadata=metadata or {},
                embedding=embedding,
                chunk_index=i,
                total_chunks=len(chunks),
                created_at=datetime.now(timezone.utc)
            )

            documents.append(doc)

        # Store in vector store
        if course_id not in self.vector_store:
            self.vector_store[course_id] = []

        self.vector_store[course_id].extend(documents)

        # Also store in Firestore for persistence
        db = get_firestore_service()
        for doc in documents:
            await db.create_subcollection_document(
                "courses", course_id, "rag_documents",
                doc.model_dump(exclude={'embedding'})  # Store embedding separately
            )

            # Store embedding in a separate collection (Firestore has size limits)
            await db.create_document(
                "embeddings",
                {"doc_id": doc.id, "embedding": doc.embedding},
                doc_id=f"{course_id}_{doc.id}"
            )

        logger.info(f"Processed {len(documents)} chunks from {file_path}")
        return documents

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF."""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document."""
        doc = docx.Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_markdown_text(self, file_path: str) -> str:
        """Extract text from Markdown."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        # Convert markdown to plain text
        html = markdown.markdown(md_content)
        # Simple HTML tag removal (in production, use BeautifulSoup)
        import re
        text = re.sub('<[^<]+?>', '', html)
        return text

    async def retrieve_context(
        self,
        query: str,
        course_id: str,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Retrieve relevant documents for a query."""

        if course_id not in self.vector_store or not self.vector_store[course_id]:
            # Try loading from Firestore
            await self._load_course_documents(course_id)

            if course_id not in self.vector_store:
                return []

        # Generate query embedding
        query_embedding = await self._generate_embedding(query)

        # Calculate similarities
        similarities = []
        for doc in self.vector_store[course_id]:
            # Apply filters
            if filters:
                if 'student_id' in filters and doc.student_id != filters['student_id']:
                    continue
                if 'assignment_id' in filters and doc.assignment_id != filters['assignment_id']:
                    continue
                if 'type' in filters and doc.type != filters['type']:
                    continue

            similarity = self._calculate_similarity(query_embedding, doc.embedding)
            if similarity >= self.config.similarity_threshold:
                similarities.append((similarity, doc))

        # Sort by similarity and return top-k
        similarities.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in similarities[:self.config.top_k]]

    async def _load_course_documents(self, course_id: str):
        """Load documents from Firestore into memory."""
        db = get_firestore_service()

        # Load documents
        docs = await db.list_subcollection(
            "courses", course_id, "rag_documents", Document
        )

        if not docs:
            return

        # Load embeddings
        for doc in docs:
            embedding_doc = await db.get_document(
                "embeddings",
                f"{course_id}_{doc.id}",
                dict
            )
            if embedding_doc:
                doc.embedding = embedding_doc.get('embedding', [])

        self.vector_store[course_id] = docs
        logger.info(f"Loaded {len(docs)} documents for course {course_id}")

    def build_contextualized_prompt(
        self,
        base_prompt: str,
        retrieved_docs: List[Document],
        max_context_tokens: int = 2000
    ) -> str:
        """Build a prompt with retrieved context."""

        if not retrieved_docs:
            return base_prompt

        context_parts = []
        current_tokens = 0

        for doc in retrieved_docs:
            doc_context = f"[{doc.type.upper()}: {doc.title}]\n{doc.content}\n"
            doc_tokens = self._count_tokens(doc_context)

            if current_tokens + doc_tokens > max_context_tokens:
                break

            context_parts.append(doc_context)
            current_tokens += doc_tokens

        if not context_parts:
            return base_prompt

        context = "\n---\n".join(context_parts)

        return f"""Based on the following context from the student's materials:

{context}

---

{base_prompt}"""

    async def generate_personalized_questions(
        self,
        student_id: str,
        course_id: str,
        assignment_id: str,
        topic: str,
        num_questions: int = 3
    ) -> List[str]:
        """Generate personalized questions based on student's submissions."""

        # Retrieve relevant context
        context_docs = await self.retrieve_context(
            query=topic,
            course_id=course_id,
            filters={
                'student_id': student_id,
                'assignment_id': assignment_id
            }
        )

        if not context_docs:
            # Fallback to generic questions
            return self._generate_generic_questions(topic, num_questions)

        # Build prompt with context
        prompt = f"""Generate {num_questions} specific, probing questions about '{topic}'
        based on the student's submitted work. The questions should:
        1. Reference specific claims or decisions from their work
        2. Test deep understanding, not memorization
        3. Be answerable in 2-3 sentences each
        4. Build on each other progressively

        Student's work excerpts:
        {' '.join([doc.content[:200] for doc in context_docs[:3]])}

        Generate the questions as a JSON array of strings."""

        # Use Gemini to generate questions
        model = genai.GenerativeModel('gemini-1.5-pro')
        response = await model.generate_content_async(prompt)

        try:
            import json
            questions = json.loads(response.text)
            return questions[:num_questions]
        except:
            return self._generate_generic_questions(topic, num_questions)

    def _generate_generic_questions(self, topic: str, num_questions: int) -> List[str]:
        """Generate generic fallback questions."""
        generic = [
            f"Can you explain the key concepts behind {topic}?",
            f"What are the main challenges when implementing {topic}?",
            f"How would you evaluate the effectiveness of {topic} in practice?",
            f"What alternatives did you consider for {topic}, and why did you choose your approach?",
            f"How does {topic} relate to the broader goals of your project?"
        ]
        return generic[:num_questions]

    async def analyze_submission_for_grading(
        self,
        student_id: str,
        course_id: str,
        assignment_id: str,
        transcript: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """Analyze student submission against their materials for grading."""

        # Get student's submitted materials
        submitted_docs = await self.retrieve_context(
            query=' '.join([msg['content'] for msg in transcript if msg['role'] == 'user'][:3]),
            course_id=course_id,
            filters={'student_id': student_id}
        )

        analysis = {
            'consistency_score': 0.0,
            'depth_demonstrated': 0.0,
            'material_coverage': 0.0,
            'specific_references': [],
            'gaps_identified': []
        }

        if not submitted_docs:
            return analysis

        # Check consistency between oral responses and written work
        for msg in transcript:
            if msg['role'] == 'user':
                response_embedding = await self._generate_embedding(msg['content'])

                for doc in submitted_docs:
                    similarity = self._calculate_similarity(response_embedding, doc.embedding)
                    if similarity > 0.8:
                        analysis['specific_references'].append({
                            'response': msg['content'][:100],
                            'source': f"{doc.type}: {doc.title}",
                            'similarity': similarity
                        })

        # Calculate scores
        analysis['consistency_score'] = min(1.0, len(analysis['specific_references']) / 5)

        if analysis['specific_references']:
            if NUMPY_AVAILABLE:
                analysis['depth_demonstrated'] = np.mean([
                    ref['similarity'] for ref in analysis['specific_references']
                ])
            else:
                # Fallback to pure Python mean calculation
                similarities = [ref['similarity'] for ref in analysis['specific_references']]
                analysis['depth_demonstrated'] = sum(similarities) / len(similarities) if similarities else 0.0
        else:
            analysis['depth_demonstrated'] = 0.0

        # Check material coverage
        doc_types_referenced = set([ref['source'].split(':')[0] for ref in analysis['specific_references']])
        doc_types_available = set([doc.type for doc in submitted_docs])
        analysis['material_coverage'] = len(doc_types_referenced) / max(1, len(doc_types_available))

        return analysis


# Singleton instance
_rag_handler: Optional[RAGHandler] = None

def get_rag_handler() -> RAGHandler:
    """Get or create RAG handler instance."""
    global _rag_handler
    if _rag_handler is None:
        _rag_handler = RAGHandler()
    return _rag_handler